"""Document processing for multiple file formats."""
import os
import mimetypes
from typing import List, Tuple
from pathlib import Path

# PDF processing
try:
    from PyPDF2 import PdfReader
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Word document processing
try:
    from docx import Document as DocxDocument
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# Markdown and HTML
import html2text
import markdown

# CSV processing
import csv
import io

class DocumentProcessor:
    """Process various document formats and extract text."""
    
    @staticmethod
    def get_file_type(filename: str) -> str:
        """Determine file type from extension."""
        ext = Path(filename).suffix.lower()
        type_map = {
            '.pdf': 'pdf',
            '.docx': 'docx',
            '.doc': 'doc',
            '.txt': 'txt',
            '.md': 'markdown',
            '.html': 'html',
            '.htm': 'html',
            '.csv': 'csv',
            '.rtf': 'rtf'
        }
        return type_map.get(ext, 'unknown')
    
    @staticmethod
    def extract_text_from_pdf(file_path: str) -> str:
        """Extract text from PDF file."""
        if not PDF_AVAILABLE:
            raise ImportError("PyPDF2 is not installed. Install with: pip install PyPDF2")
        
        text = ""
        try:
            # Try pdfplumber first for better extraction
            if PDFPLUMBER_AVAILABLE:
                with pdfplumber.open(file_path) as pdf:
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
            else:
                # Fallback to PyPDF2
                with open(file_path, "rb") as f:
                    reader = PdfReader(f)
                    for page in reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n\n"
        except Exception as e:
            raise Exception(f"Error extracting text from PDF: {str(e)}")
        
        return text.strip()
    
    @staticmethod
    def extract_text_from_docx(file_path: str) -> str:
        """Extract text from Word document."""
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is not installed. Install with: pip install python-docx")
        
        try:
            doc = DocxDocument(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            
            return "\n\n".join(paragraphs)
        except Exception as e:
            raise Exception(f"Error extracting text from Word document: {str(e)}")
    
    @staticmethod
    def extract_text_from_txt(file_path: str) -> str:
        """Extract text from plain text file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with different encoding
            with open(file_path, 'r', encoding='latin-1', errors='ignore') as f:
                return f.read()
    
    @staticmethod
    def extract_text_from_markdown(file_path: str) -> str:
        """Extract text from Markdown file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                md_text = f.read()
            # Convert markdown to plain text
            html = markdown.markdown(md_text)
            h = html2text.HTML2Text()
            h.ignore_links = False
            return h.handle(html)
        except Exception as e:
            raise Exception(f"Error extracting text from Markdown: {str(e)}")
    
    @staticmethod
    def extract_text_from_html(file_path: str) -> str:
        """Extract text from HTML file."""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            h = html2text.HTML2Text()
            h.ignore_links = False
            return h.handle(html_content)
        except Exception as e:
            raise Exception(f"Error extracting text from HTML: {str(e)}")
    
    @staticmethod
    def extract_text_from_csv(file_path: str) -> str:
        """Extract text from CSV file."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                reader = csv.reader(f)
                rows = []
                for row in reader:
                    rows.append(" | ".join(row))
                return "\n".join(rows)
        except Exception as e:
            raise Exception(f"Error extracting text from CSV: {str(e)}")
    
    @staticmethod
    def extract_text_from_rtf(file_path: str) -> str:
        """Extract text from RTF file (basic extraction)."""
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            # Basic RTF text extraction (remove RTF control words)
            import re
            # Remove RTF control words and braces
            text = re.sub(r'\{[^}]*\}', '', content)
            text = re.sub(r'\\[a-z]+\d*\s?', '', text)
            return text.strip()
        except Exception as e:
            raise Exception(f"Error extracting text from RTF: {str(e)}")
    
    @classmethod
    def process_document(cls, file_path: str, filename: str) -> Tuple[str, str]:
        """Process a document and extract text based on file type.
        
        Returns:
            Tuple of (extracted_text, file_type)
        """
        file_type = cls.get_file_type(filename)
        
        if file_type == 'pdf':
            text = cls.extract_text_from_pdf(file_path)
        elif file_type == 'docx':
            text = cls.extract_text_from_docx(file_path)
        elif file_type == 'txt':
            text = cls.extract_text_from_txt(file_path)
        elif file_type == 'markdown':
            text = cls.extract_text_from_markdown(file_path)
        elif file_type == 'html':
            text = cls.extract_text_from_html(file_path)
        elif file_type == 'csv':
            text = cls.extract_text_from_csv(file_path)
        elif file_type == 'rtf':
            text = cls.extract_text_from_rtf(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
        
        if not text or not text.strip():
            raise ValueError(f"No text could be extracted from {filename}")
        
        return text.strip(), file_type

